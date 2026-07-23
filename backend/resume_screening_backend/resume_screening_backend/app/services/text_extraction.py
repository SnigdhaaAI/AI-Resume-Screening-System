"""
Extracts raw text from uploaded resume/job files (.pdf, .docx, .doc, .txt).
"""
import os
import pdfplumber
from docx import Document
from fastapi import HTTPException


def extract_text_from_pdf(file_path: str) -> str:
    text_chunks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Failed to parse PDF: {exc}")
    return "\n".join(text_chunks)


def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text out of tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Failed to parse DOCX: {exc}")


def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Failed to read TXT file: {exc}")


def extract_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text = extract_text_from_docx(file_path)
    elif ext == ".txt":
        text = extract_text_from_txt(file_path)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    if not text or not text.strip():
        raise HTTPException(
            status_code=422,
            detail="No extractable text found in file (it may be a scanned/image-based document).",
        )
    return text.strip()
